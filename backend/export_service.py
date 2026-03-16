from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import openpyxl
from io import BytesIO
from typing import List
import models
import datetime

class ExportService:
    @staticmethod
    def to_docx(logs: List[models.FaultLog]) -> BytesIO:
        doc = Document()
        doc.add_heading('Fault Logs Report', 0)
        
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Company'
        hdr_cells[2].text = 'Type'
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Created'
        hdr_cells[5].text = 'Logged By'
        
        for log in logs:
            row_cells = table.add_row().cells
            row_cells[0].text = str(log.id)
            row_cells[1].text = log.company_name
            row_cells[2].text = log.fault_type
            row_cells[3].text = log.status
            row_cells[4].text = log.created_at.strftime("%Y-%m-%d %H:%M")
            row_cells[5].text = log.logged_by
            
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def to_pdf(logs: List[models.FaultLog]) -> BytesIO:
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.drawString(100, 750, "Fault Logs Report")
        y = 730
        for log in logs:
            p.drawString(100, y, f"ID: {log.id} | {log.company_name} | {log.fault_type} | {log.status}")
            y -= 20
            if y < 50:
                p.showPage()
                y = 750
        p.save()
        buffer.seek(0)
        return buffer

    @staticmethod
    def to_excel(logs: List[models.FaultLog]) -> BytesIO:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Fault Logs"
        
        headers = ['ID', 'ISP', 'Location', 'Type', 'Severity', 'Status', 'Created', 'Resolved']
        ws.append(headers)
        
        for log in logs:
            created_str = log.created_at.strftime("%Y-%m-%d %H:%M") if log.created_at else ""
            resolved_str = log.resolved_at.strftime("%Y-%m-%d %H:%M") if log.resolved_at else ""
            ws.append([
                log.id,
                log.isp_name if hasattr(log, 'isp_name') else log.company_name,
                getattr(log, 'location', ''),
                log.fault_type,
                getattr(log, 'severity', ''),
                log.status,
                created_str,
                resolved_str
            ])
            
        # Adjust column widths so times are visible
        ws.column_dimensions['G'].width = 18
        ws.column_dimensions['H'].width = 18
            
        file_stream = BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def isp_report_docx(chart_data: dict) -> BytesIO:
        """Generate a formatted Word document ISP performance report."""
        doc = Document()

        # --- Title ---
        title = doc.add_heading('ISP Performance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(
            f'Bulawayo City Council — Network Infrastructure Management System\n'
            f'Generated: {datetime.datetime.now().strftime("%d %B %Y, %H:%M")}'
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

        # --- Summary section ---
        doc.add_heading('Executive Summary', level=1)
        isps = chart_data.get('isps', [])
        summary_parts = []
        for isp in isps:
            rate = isp.get('resolution_rate', 0)
            summary_parts.append(
                f"{isp['name']} logged {isp['total_faults']} fault(s) with a resolution rate of {rate:.0f}%"
            )
        summary_text = '. '.join(summary_parts) + '.' if summary_parts else 'No data available.'
        doc.add_paragraph(summary_text)

        doc.add_paragraph()

        # --- Performance Table ---
        doc.add_heading('ISP Performance Breakdown', level=1)

        headers = ['ISP Provider', 'Total Faults', 'Resolved', 'Critical', 'Resolution Rate']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row — bold, dark blue background
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Apply shading via XML
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0F3460')
            tc_pr.append(shd)

        # Data rows
        for isp in isps:
            rate = isp.get('resolution_rate', 0)
            row_data = [
                isp.get('name', ''),
                str(isp.get('total_faults', 0)),
                str(isp.get('resolved', 0)),
                str(isp.get('critical', 0)),
                f"{rate:.0f}%"
            ]
            row = table.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = val
                # Colour-code the resolution rate column
                if i == 4:
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True
                    if rate >= 80:
                        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)  # green
                    elif rate >= 50:
                        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)  # amber
                    else:
                        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red

        doc.add_paragraph()

        # --- Footer note ---
        footer = doc.add_paragraph(
            'This report was auto-generated by BCC NIMS. '
            'Data reflects real-time fault log records.'
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

